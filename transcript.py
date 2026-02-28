import json
import ollama

MODEL = "llama3.1:8b"

PARSE_PROMPT = """You are a transcript parser. Extract every course from the transcript text below into a JSON array.

Each course object must have these fields:
- "name": course name (string)
- "grade": letter grade normalized to one of: A+, A, A-, B+, B, B-, C+, C, C-, D+, D, D-, F
- "year": one of "Freshman", "Sophomore", "Junior", "Senior" (infer from grade level, year labels, or 9th/10th/11th/12th)
- "course_type": one of "Regular", "Honors", "AP", "IB", "Dual Enrollment" (infer from course name or labels)
- "credits": number of credits (default 1.0 if not listed)

Rules:
- Convert percentage grades to letter grades (90-100=A, 80-89=B, etc.)
- If a course name contains "AP " or "Advanced Placement", set course_type to "AP"
- If a course name contains "IB " or "International Baccalaureate", set course_type to "IB"
- If a course name contains "Honors" or "Hon ", set course_type to "Honors"
- If a course name contains "Dual Enrollment" or "DE " or "College ", set course_type to "Dual Enrollment"
- Otherwise set course_type to "Regular"
- If year/grade level is unclear, use "Junior" as default

Respond with ONLY a valid JSON array. No markdown, no explanation."""


def extract_text_pdf(file_path):
    """Extract text from a PDF file using pdfplumber."""
    import pdfplumber

    text_parts = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def extract_text_docx(file_path):
    """Extract text from a DOCX file using python-docx."""
    import docx

    doc = docx.Document(file_path)
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)
    return "\n".join(text_parts)


def extract_text(file_path, filename):
    """Extract text from a file based on its extension."""
    ext = filename.rsplit(".", 1)[-1].lower()
    if ext == "pdf":
        return extract_text_pdf(file_path)
    elif ext in ("docx", "doc"):
        return extract_text_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: .{ext}")


def parse_transcript(raw_text):
    """Send extracted text to Ollama and parse into structured course data."""
    response = ollama.chat(
        model=MODEL,
        messages=[
            {"role": "system", "content": PARSE_PROMPT},
            {"role": "user", "content": f"Transcript text:\n\n{raw_text}"},
        ],
    )

    text = response["message"]["content"].strip()

    # Extract JSON array from response
    if text.startswith("```"):
        lines = text.split("\n")
        text = "\n".join(lines[1:-1])
    if not text.startswith("["):
        start = text.index("[")
        end = text.rindex("]") + 1
        text = text[start:end]

    courses = json.loads(text)

    # Validate and normalize
    valid_grades = {"A+", "A", "A-", "B+", "B", "B-", "C+", "C", "C-", "D+", "D", "D-", "F"}
    valid_years = {"Freshman", "Sophomore", "Junior", "Senior"}
    valid_types = {"Regular", "Honors", "AP", "IB", "Dual Enrollment"}

    cleaned = []
    for c in courses:
        if not isinstance(c, dict) or "name" not in c:
            continue
        grade = c.get("grade", "B")
        if grade not in valid_grades:
            grade = "B"
        year = c.get("year", "Junior")
        if year not in valid_years:
            year = "Junior"
        ctype = c.get("course_type", "Regular")
        if ctype not in valid_types:
            ctype = "Regular"
        try:
            credits = float(c.get("credits", 1.0))
        except (ValueError, TypeError):
            credits = 1.0

        cleaned.append({
            "name": str(c["name"]).strip(),
            "grade": grade,
            "year": year,
            "course_type": ctype,
            "credits": credits,
        })

    return cleaned
